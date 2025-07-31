"""File handling utilities for document processing."""

import logging
from pathlib import Path
from typing import Union, Optional
import tempfile
import hashlib
from datetime import datetime

import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
import streamlit as st

from config.settings import settings


logger = logging.getLogger(__name__)


class FileHandler:
    """Handle file operations for document processing."""
    
    def __init__(self):
        self.upload_dir = settings.upload_dir
        self.supported_types = settings.supported_file_types
        self._ensure_directories()
    
    def _ensure_directories(self):
        """Ensure required directories exist."""
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    def validate_file(self, file) -> bool:
        """Validate uploaded file."""
        # Check file size
        file_size = file.size if hasattr(file, 'size') else len(file.getvalue())
        if file_size > settings.max_upload_size_mb * 1024 * 1024:
            raise ValueError(f"File size exceeds {settings.max_upload_size_mb}MB limit")
        
        # Check file type
        file_ext = Path(file.name).suffix.lower().lstrip('.')
        if file_ext not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        return True
    
    def save_uploaded_file(self, uploaded_file) -> Path:
        """Save uploaded file to temporary location."""
        try:
            # Validate file
            self.validate_file(uploaded_file)
            
            # Generate unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_hash = hashlib.md5(uploaded_file.name.encode()).hexdigest()[:8]
            file_ext = Path(uploaded_file.name).suffix
            filename = f"{timestamp}_{file_hash}{file_ext}"
            
            # Save file
            file_path = self.upload_dir / filename
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            logger.info(f"File saved: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Failed to save file: {e}")
            raise
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text content from various file types."""
        suffix = file_path.suffix.lower()
        
        try:
            if suffix == '.txt':
                return self._extract_from_txt(file_path)
            elif suffix == '.pdf':
                return self._extract_from_pdf(file_path)
            elif suffix == '.docx':
                return self._extract_from_docx(file_path)
            elif suffix in ['.html', '.htm']:
                return self._extract_from_html(file_path)
            else:
                raise ValueError(f"Unsupported file type: {suffix}")
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            raise
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Unable to decode text file with common encodings")
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num + 1}: {e}")
        
        if not text_content:
            raise ValueError("No text content found in PDF")
        
        return "\n\n".join(text_content)
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    table_data.append(" | ".join(row_data))
            if table_data:
                tables_text.append("\n".join(table_data))
        
        # Combine all text
        all_text = paragraphs + tables_text
        
        if not all_text:
            raise ValueError("No text content found in DOCX")
        
        return "\n\n".join(all_text)
    
    def _extract_from_html(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        html_content = file_path.read_text(encoding='utf-8')
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for element in soup(['script', 'style']):
            element.decompose()
        
        # Extract text
        text = soup.get_text(separator='\n', strip=True)
        
        if not text:
            raise ValueError("No text content found in HTML")
        
        return text
    
    def get_file_metadata(self, file_path: Path) -> dict:
        """Extract metadata from file."""
        stat = file_path.stat()
        
        metadata = {
            "filename": file_path.name,
            "size_bytes": stat.st_size,
            "size_mb": round(stat.st_size / (1024 * 1024), 2),
            "created": datetime.fromtimestamp(stat.st_ctime),
            "modified": datetime.fromtimestamp(stat.st_mtime),
            "type": file_path.suffix.lower().lstrip('.')
        }
        
        # PDF-specific metadata
        if file_path.suffix.lower() == '.pdf':
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata["pages"] = len(pdf_reader.pages)
                    
                    if pdf_reader.metadata:
                        metadata["pdf_metadata"] = {
                            "title": pdf_reader.metadata.get('/Title', ''),
                            "author": pdf_reader.metadata.get('/Author', ''),
                            "subject": pdf_reader.metadata.get('/Subject', ''),
                            "creator": pdf_reader.metadata.get('/Creator', '')
                        }
            except Exception as e:
                logger.warning(f"Failed to extract PDF metadata: {e}")
        
        return metadata
    
    def cleanup_old_files(self, hours: int = 24):
        """Remove files older than specified hours."""
        current_time = datetime.now()
        removed_count = 0
        
        for file_path in self.upload_dir.glob("*"):
            if file_path.is_file():
                file_age = current_time - datetime.fromtimestamp(file_path.stat().st_mtime)
                if file_age.total_seconds() > hours * 3600:
                    try:
                        file_path.unlink()
                        removed_count += 1
                    except Exception as e:
                        logger.error(f"Failed to remove old file {file_path}: {e}")
        
        if removed_count > 0:
            logger.info(f"Cleaned up {removed_count} old files")
        
        return removed_count