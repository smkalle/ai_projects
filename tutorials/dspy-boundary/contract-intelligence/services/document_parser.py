"""Document parsing service for various file formats."""

import io
import os
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Union, Any
from dataclasses import dataclass
from datetime import datetime

import fitz  # PyMuPDF
from docx import Document
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes
import magic

from config.settings import settings


@dataclass
class DocumentMetadata:
    """Document metadata structure."""
    filename: str
    file_size: int
    file_type: str
    mime_type: str
    page_count: int
    created_date: Optional[datetime] = None
    modified_date: Optional[datetime] = None
    author: Optional[str] = None
    title: Optional[str] = None
    subject: Optional[str] = None
    language: Optional[str] = None


@dataclass
class ParsedDocument:
    """Parsed document structure."""
    content: str
    metadata: DocumentMetadata
    pages: List[str]
    images: List[str] = None
    tables: List[Dict] = None
    confidence_score: float = 1.0
    parsing_method: str = "direct"


class DocumentParsingError(Exception):
    """Custom exception for document parsing errors."""
    pass


class DocumentParser:
    """Service for parsing various document formats."""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._parse_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._parse_docx,
            'application/msword': self._parse_doc,
            'text/plain': self._parse_text,
        }
    
    def parse_document(self, file_data: bytes, filename: str) -> ParsedDocument:
        """
        Parse document from bytes data.
        
        Args:
            file_data: Raw file bytes
            filename: Original filename
            
        Returns:
            ParsedDocument with extracted content and metadata
            
        Raises:
            DocumentParsingError: If parsing fails
        """
        try:
            # Detect file type
            mime_type = magic.from_buffer(file_data, mime=True)
            
            # Validate file size
            if len(file_data) > settings.get_max_file_size_bytes():
                raise DocumentParsingError(
                    f"File size {len(file_data)} exceeds limit {settings.get_max_file_size_bytes()}"
                )
            
            # Check supported format
            if mime_type not in self.supported_types:
                raise DocumentParsingError(f"Unsupported file type: {mime_type}")
            
            # Parse based on type
            parser_func = self.supported_types[mime_type]
            result = parser_func(file_data, filename)
            
            # Validate page count
            if result.metadata.page_count > settings.max_pages_per_document:
                raise DocumentParsingError(
                    f"Document has {result.metadata.page_count} pages, "
                    f"exceeds limit of {settings.max_pages_per_document}"
                )
            
            return result
            
        except Exception as e:
            if isinstance(e, DocumentParsingError):
                raise
            raise DocumentParsingError(f"Failed to parse document: {str(e)}")
    
    def _parse_pdf(self, file_data: bytes, filename: str) -> ParsedDocument:
        """Parse PDF document."""
        try:
            # Open PDF with PyMuPDF
            doc = fitz.open(stream=file_data, filetype="pdf")
            
            pages = []
            full_content = []
            images = []
            confidence_score = 1.0
            parsing_method = "direct"
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extract text
                text = page.get_text()
                pages.append(text)
                full_content.append(text)
                
                # If no text found, try OCR
                if not text.strip() and settings.enable_ocr:
                    try:
                        # Convert page to image
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        image = Image.open(io.BytesIO(img_data))
                        
                        # OCR the image
                        ocr_text = pytesseract.image_to_string(image)
                        pages[-1] = ocr_text
                        full_content[-1] = ocr_text
                        
                        # Lower confidence for OCR
                        confidence_score = min(confidence_score, 0.8)
                        parsing_method = "ocr"
                        
                    except Exception as ocr_error:
                        pages[-1] = f"[OCR Error: {str(ocr_error)}]"
                        confidence_score = min(confidence_score, 0.5)
                
                # Extract images
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_ext = base_image["ext"]
                        image_bytes = base_image["image"]
                        
                        # Store image reference
                        images.append(f"page_{page_num}_image_{img_index}.{image_ext}")
                    except Exception:
                        continue
            
            # Extract metadata
            pdf_metadata = doc.metadata
            metadata = DocumentMetadata(
                filename=filename,
                file_size=len(file_data),
                file_type="pdf",
                mime_type="application/pdf",
                page_count=len(doc),
                title=pdf_metadata.get('title'),
                author=pdf_metadata.get('author'),
                subject=pdf_metadata.get('subject'),
                created_date=self._parse_pdf_date(pdf_metadata.get('creationDate')),
                modified_date=self._parse_pdf_date(pdf_metadata.get('modDate'))
            )
            
            doc.close()
            
            return ParsedDocument(
                content="\n\n".join(full_content),
                metadata=metadata,
                pages=pages,
                images=images,
                confidence_score=confidence_score,
                parsing_method=parsing_method
            )
            
        except Exception as e:
            raise DocumentParsingError(f"PDF parsing failed: {str(e)}")
    
    def _parse_docx(self, file_data: bytes, filename: str) -> ParsedDocument:
        """Parse DOCX document."""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(file_data)
                tmp_path = tmp.name
            
            try:
                # Open document
                doc = Document(tmp_path)
                
                # Extract text by paragraphs
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                content = "\n".join(paragraphs)
                
                # Extract metadata from core properties
                props = doc.core_properties
                metadata = DocumentMetadata(
                    filename=filename,
                    file_size=len(file_data),
                    file_type="docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    page_count=1,  # DOCX doesn't have clear page breaks
                    title=props.title,
                    author=props.author,
                    subject=props.subject,
                    created_date=props.created,
                    modified_date=props.modified
                )
                
                return ParsedDocument(
                    content=content,
                    metadata=metadata,
                    pages=[content],  # Single page for DOCX
                    confidence_score=1.0,
                    parsing_method="direct"
                )
                
            finally:
                # Clean up temporary file
                os.unlink(tmp_path)
                
        except Exception as e:
            raise DocumentParsingError(f"DOCX parsing failed: {str(e)}")
    
    def _parse_doc(self, file_data: bytes, filename: str) -> ParsedDocument:
        """Parse legacy DOC document."""
        # For now, treat as text extraction challenge
        # In production, you might use python-docx2txt or antiword
        raise DocumentParsingError("Legacy DOC format not supported yet. Please convert to DOCX.")
    
    def _parse_text(self, file_data: bytes, filename: str) -> ParsedDocument:
        """Parse plain text document."""
        try:
            # Decode text
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            content = None
            
            for encoding in encodings:
                try:
                    content = file_data.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise DocumentParsingError("Unable to decode text file")
            
            # Estimate page count (rough approximation)
            page_count = max(1, len(content) // 3000)
            
            metadata = DocumentMetadata(
                filename=filename,
                file_size=len(file_data),
                file_type="txt",
                mime_type="text/plain",
                page_count=page_count,
                modified_date=datetime.now()
            )
            
            return ParsedDocument(
                content=content,
                metadata=metadata,
                pages=[content],
                confidence_score=1.0,
                parsing_method="direct"
            )
            
        except Exception as e:
            raise DocumentParsingError(f"Text parsing failed: {str(e)}")
    
    def _parse_pdf_date(self, date_str: Optional[str]) -> Optional[datetime]:
        """Parse PDF date string to datetime."""
        if not date_str:
            return None
        
        try:
            # PDF dates are in format: D:YYYYMMDDHHmmSSOHH'mm'
            if date_str.startswith("D:"):
                date_str = date_str[2:]
            
            # Extract just the date part
            date_part = date_str[:14]
            return datetime.strptime(date_part, "%Y%m%d%H%M%S")
        except Exception:
            return None
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return ['.pdf', '.docx', '.doc', '.txt']
    
    def validate_file(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """
        Validate file before parsing.
        
        Returns:
            Dictionary with validation results
        """
        result = {
            'valid': True,
            'errors': [],
            'warnings': [],
            'file_info': {}
        }
        
        try:
            # Check file size
            if len(file_data) > settings.get_max_file_size_bytes():
                result['valid'] = False
                result['errors'].append(
                    f"File size {len(file_data):,} bytes exceeds limit "
                    f"{settings.get_max_file_size_bytes():,} bytes"
                )
            
            # Detect file type
            mime_type = magic.from_buffer(file_data, mime=True)
            result['file_info']['mime_type'] = mime_type
            result['file_info']['detected_extension'] = self._mime_to_extension(mime_type)
            
            # Check supported format
            if mime_type not in self.supported_types:
                result['valid'] = False
                result['errors'].append(f"Unsupported file type: {mime_type}")
            
            # Check file extension matches content
            file_ext = Path(filename).suffix.lower()
            expected_ext = self._mime_to_extension(mime_type)
            if file_ext != expected_ext:
                result['warnings'].append(
                    f"File extension {file_ext} doesn't match detected type {expected_ext}"
                )
            
            # Check if file is empty
            if len(file_data) == 0:
                result['valid'] = False
                result['errors'].append("File is empty")
            
        except Exception as e:
            result['valid'] = False
            result['errors'].append(f"Validation error: {str(e)}")
        
        return result
    
    def _mime_to_extension(self, mime_type: str) -> str:
        """Convert MIME type to file extension."""
        mime_map = {
            'application/pdf': '.pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'application/msword': '.doc',
            'text/plain': '.txt'
        }
        return mime_map.get(mime_type, '')


# Global parser instance
document_parser = DocumentParser()