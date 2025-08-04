"""Contract parser service for extracting text and metadata from documents."""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import logging

# Document parsing libraries
try:
    import pymupdf as fitz  # PyMuPDF for PDF parsing
except ImportError:
    import fitz

from docx import Document
import pytesseract
from PIL import Image
import io

from models.document import DocumentMetadata, ParsedDocument, DocumentChunk
from config.settings import settings

logger = logging.getLogger(__name__)


class ContractParser:
    """Service for parsing contracts from various file formats."""
    
    def __init__(self):
        self.max_chunk_size = 2000  # Characters per chunk
        self.chunk_overlap = 200    # Overlap between chunks
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt'}
    
    def parse_document(self, file_path: str) -> ParsedDocument:
        """Parse a document and extract text with metadata."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Extract metadata
        metadata = self._extract_metadata(file_path)
        
        # Parse based on file type
        if file_extension == '.pdf':
            text, page_texts = self._parse_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            text, page_texts = self._parse_docx(file_path)
        elif file_extension == '.txt':
            text, page_texts = self._parse_txt(file_path)
        else:
            raise ValueError(f"Parser not implemented for: {file_extension}")
        
        # Clean and normalize text
        text = self._clean_text(text)
        
        # Create chunks for processing
        chunks = self._create_chunks(text, page_texts)
        
        return ParsedDocument(
            filename=file_path.name,
            file_path=str(file_path),
            text=text,
            chunks=chunks,
            metadata=metadata,
            page_count=len(page_texts),
            word_count=len(text.split()),
            char_count=len(text)
        )
    
    def _extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract file metadata."""
        stat = file_path.stat()
        
        return DocumentMetadata(
            file_size=stat.st_size,
            file_type=file_path.suffix.lower(),
            created_at=stat.st_ctime,
            modified_at=stat.st_mtime,
            mime_type=self._get_mime_type(file_path.suffix.lower())
        )
    
    def _parse_pdf(self, file_path: Path) -> Tuple[str, List[str]]:
        """Parse PDF file and extract text."""
        try:
            doc = fitz.open(str(file_path))
            full_text = []
            page_texts = []
            
            for page_num, page in enumerate(doc):
                # Extract text
                page_text = page.get_text()
                
                # If no text found, try OCR
                if not page_text.strip() and settings.enable_ocr:
                    page_text = self._ocr_page(page)
                
                page_texts.append(page_text)
                full_text.append(f"\n--- Page {page_num + 1} ---\n{page_text}")
            
            doc.close()
            return '\n'.join(full_text), page_texts
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            raise
    
    def _parse_docx(self, file_path: Path) -> Tuple[str, List[str]]:
        """Parse DOCX file and extract text."""
        try:
            doc = Document(str(file_path))
            full_text = []
            page_texts = []
            current_page = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    current_page.append(text)
                    
                    # Simple page break detection
                    if len('\n'.join(current_page)) > 3000:
                        page_texts.append('\n'.join(current_page))
                        current_page = []
            
            # Add remaining text
            if current_page:
                page_texts.append('\n'.join(current_page))
            
            # If no pages detected, treat as single page
            if not page_texts:
                page_texts = ['\n'.join(p.text for p in doc.paragraphs if p.text.strip())]
            
            full_text = '\n\n'.join(page_texts)
            return full_text, page_texts
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
    
    def _parse_txt(self, file_path: Path) -> Tuple[str, List[str]]:
        """Parse TXT file and extract text."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Simple pagination for text files
            lines = text.split('\n')
            pages = []
            current_page = []
            line_count = 0
            
            for line in lines:
                current_page.append(line)
                line_count += 1
                
                if line_count >= 50:  # ~50 lines per page
                    pages.append('\n'.join(current_page))
                    current_page = []
                    line_count = 0
            
            if current_page:
                pages.append('\n'.join(current_page))
            
            return text, pages if pages else [text]
            
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {str(e)}")
            raise
    
    def _ocr_page(self, page) -> str:
        """Perform OCR on a PDF page."""
        try:
            # Render page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x scale for better OCR
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # Perform OCR
            text = pytesseract.image_to_string(img)
            return text
            
        except Exception as e:
            logger.warning(f"OCR failed: {str(e)}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', text)
        
        # Normalize quotes and dashes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        text = text.replace('–', '-').replace('—', '-')
        
        # Remove page numbers and headers/footers patterns
        text = re.sub(r'Page \d+ of \d+', '', text)
        text = re.sub(r'\n\d+\n', '\n', text)
        
        return text.strip()
    
    def _create_chunks(self, text: str, page_texts: List[str]) -> List[DocumentChunk]:
        """Create overlapping chunks from text."""
        chunks = []
        
        # Create chunks from full text
        words = text.split()
        current_chunk = []
        current_size = 0
        chunk_id = 0
        
        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1
            
            if current_size >= self.max_chunk_size:
                # Create chunk
                chunk_text = ' '.join(current_chunk)
                chunks.append(DocumentChunk(
                    id=f"chunk_{chunk_id}",
                    text=chunk_text,
                    start_char=max(0, text.find(chunk_text)),
                    end_char=min(len(text), text.find(chunk_text) + len(chunk_text)),
                    metadata={"chunk_id": chunk_id}
                ))
                
                # Keep overlap
                overlap_words = int(self.chunk_overlap / 5)  # Approximate words
                current_chunk = current_chunk[-overlap_words:]
                current_size = sum(len(w) + 1 for w in current_chunk)
                chunk_id += 1
        
        # Add remaining chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append(DocumentChunk(
                id=f"chunk_{chunk_id}",
                text=chunk_text,
                start_char=max(0, text.find(chunk_text)),
                end_char=min(len(text), text.find(chunk_text) + len(chunk_text)),
                metadata={"chunk_id": chunk_id}
            ))
        
        return chunks
    
    def _get_mime_type(self, extension: str) -> str:
        """Get MIME type for file extension."""
        mime_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.txt': 'text/plain'
        }
        return mime_types.get(extension, 'application/octet-stream')
    
    def extract_sections(self, parsed_doc: ParsedDocument) -> Dict[str, str]:
        """Extract common contract sections."""
        sections = {}
        text = parsed_doc.text.lower()
        
        # Common section patterns
        section_patterns = {
            'parties': r'(parties|between.*and)',
            'definitions': r'(definitions|defined terms)',
            'scope': r'(scope of work|services|deliverables)',
            'payment': r'(payment|compensation|fees)',
            'term': r'(term|duration|period)',
            'termination': r'(termination|expiration)',
            'confidentiality': r'(confidential|non-disclosure|proprietary)',
            'liability': r'(liability|limitation|damages)',
            'warranty': r'(warrant|representation|guarantee)',
            'governing_law': r'(governing law|jurisdiction|disputes)'
        }
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text)
            if match:
                # Extract text around the match
                start = max(0, match.start() - 100)
                end = min(len(text), match.end() + 1000)
                sections[section_name] = parsed_doc.text[start:end]
        
        return sections